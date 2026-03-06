"""
Document Generator Tool

Creates adversarial DOCX documents with stealth payload embedding for
authorized security testing of LLM document ingestion systems.

This tool is designed exclusively for:
- Authorized penetration testing
- Security research
- Red team exercises
- Compliance audits

All generated documents include metadata indicating their test nature.
"""

import base64
import logging
import re
from datetime import datetime
from enum import Enum
from io import BytesIO
from typing import Any, Dict, List, Optional, Union

from . import BaseTool

logger = logging.getLogger(__name__)


# Zero-width Unicode characters for stealth embedding
ZERO_WIDTH_CHARS = {
    "zero_width_space": "\u200B",
    "zero_width_non_joiner": "\u200C",
    "zero_width_joiner": "\u200D",
    "left_to_right_mark": "\u200E",
    "right_to_left_mark": "\u200F",
    "word_joiner": "\u2060",
    "invisible_plus": "\u2064",
}


class DocumentStyle(str, Enum):
    """Predefined document styles for various document types."""

    BUSINESS_MEMO = "business_memo"
    TECHNICAL_REPORT = "technical_report"
    FINANCIAL_STATEMENT = "financial_statement"
    INTERNAL_POLICY = "internal_policy"
    PROJECT_PROPOSAL = "project_proposal"
    MEETING_MINUTES = "meeting_minutes"
    EMAIL_ATTACHMENT = "email_attachment"
    RESUME = "resume"
    ACADEMIC_PAPER = "academic_paper"
    INVOICE = "invoice"


class DocumentGeneratorTool(BaseTool):
    """
    Tool for generating adversarial DOCX documents with stealth payloads.

    Supports multiple stealth techniques:
    - Font color matching (text color = background color)
    - Zero-width Unicode characters
    - White-on-white text
    - Tiny font (1pt)
    - Metadata injection
    - Hidden text property
    - Footnote embedding
    - Comment injection
    - Header/footer hidden
    - Invisible table borders
    - Overlay text
    """

    # Document templates with structure
    DOCUMENT_TEMPLATES = {
        DocumentStyle.BUSINESS_MEMO: {
            "title": "MEMORANDUM",
            "headers": ["To:", "From:", "Date:", "Subject:"],
            "structure": "memo",
        },
        DocumentStyle.TECHNICAL_REPORT: {
            "title": "TECHNICAL REPORT",
            "headers": ["Title:", "Author:", "Date:", "Document ID:"],
            "structure": "report",
        },
        DocumentStyle.FINANCIAL_STATEMENT: {
            "title": "FINANCIAL STATEMENT",
            "headers": ["Period:", "Prepared By:", "Date:"],
            "structure": "financial",
        },
        DocumentStyle.INTERNAL_POLICY: {
            "title": "INTERNAL POLICY",
            "headers": ["Policy Number:", "Effective Date:", "Department:"],
            "structure": "policy",
        },
        DocumentStyle.PROJECT_PROPOSAL: {
            "title": "PROJECT PROPOSAL",
            "headers": ["Project:", "Submitted By:", "Date:"],
            "structure": "proposal",
        },
        DocumentStyle.MEETING_MINUTES: {
            "title": "MEETING MINUTES",
            "headers": ["Meeting:", "Date:", "Attendees:", "Time:"],
            "structure": "minutes",
        },
        DocumentStyle.EMAIL_ATTACHMENT: {
            "title": "",
            "headers": [],
            "structure": "email",
        },
        DocumentStyle.RESUME: {
            "title": "",
            "headers": [],
            "structure": "resume",
        },
        DocumentStyle.ACADEMIC_PAPER: {
            "title": "",
            "headers": ["Abstract", "Introduction", "Methodology"],
            "structure": "academic",
        },
        DocumentStyle.INVOICE: {
            "title": "INVOICE",
            "headers": ["Invoice #:", "Date:", "Bill To:", "From:"],
            "structure": "invoice",
        },
    }

    def __init__(self):
        super().__init__(
            name="docx_generator",
            description="Generate adversarial DOCX documents with stealth payload embedding for security testing.",
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "visible_content": {
                    "type": "string",
                    "description": "The visible content that appears in the document",
                },
                "payload_embeddings": {
                    "type": "array",
                    "description": "List of payloads to embed with stealth techniques",
                    "items": {
                        "type": "object",
                        "properties": {
                            "payload": {
                                "type": "string",
                                "description": "The payload text to embed",
                            },
                            "technique": {
                                "type": "string",
                                "enum": [
                                    "font_color_match",
                                    "zero_width_unicode",
                                    "white_on_white",
                                    "tiny_font",
                                    "hidden_text_property",
                                    "metadata_injection",
                                    "footnote_embedding",
                                    "comment_injection",
                                    "header_footer_hidden",
                                    "invisible_table",
                                    "overlay_text",
                                ],
                                "description": "The stealth technique to use",
                            },
                            "position": {
                                "type": "string",
                                "enum": ["start", "end", "after_text", "before_text"],
                                "description": "Where to embed the payload",
                            },
                            "anchor_text": {
                                "type": "string",
                                "description": "Anchor text for position-based embedding",
                            },
                            "additional_config": {
                                "type": "object",
                                "description": "Additional configuration for the technique",
                            },
                        },
                        "required": ["payload", "technique"],
                    },
                },
                "template": {
                    "type": "string",
                    "enum": [s.value for s in DocumentStyle],
                    "description": "Document template to use",
                },
                "metadata": {
                    "type": "object",
                    "description": "Custom document metadata",
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename (without extension)",
                },
                "return_base64": {
                    "type": "boolean",
                    "description": "Return document as base64 encoded string",
                },
            },
            "required": ["visible_content"],
        }

    async def call(
        self,
        visible_content: str,
        payload_embeddings: Optional[List[Dict[str, Any]]] = None,
        template: str = "business_memo",
        metadata: Optional[Dict[str, str]] = None,
        filename: Optional[str] = None,
        return_base64: bool = True,
    ) -> Dict[str, Any]:
        """
        Generate an adversarial DOCX document.

        Args:
            visible_content: The visible content in the document
            payload_embeddings: Payloads to embed with stealth techniques
            template: Document template style
            metadata: Custom document metadata
            filename: Output filename
            return_base64: Whether to return base64 encoded content

        Returns:
            Dict with generated document info
        """
        try:
            from docx import Document
            from docx.shared import RGBColor, Pt
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "python-docx is required for document generation. "
                "Install it with: pip install python-docx"
            )

        # Create document
        doc = Document()

        # Parse template style
        try:
            style = DocumentStyle(template)
        except ValueError:
            style = DocumentStyle.BUSINESS_MEMO

        # Apply template structure
        self._apply_template(doc, style, visible_content)

        # Apply metadata
        if metadata:
            self._apply_metadata(doc, metadata)

        # Add security notice in metadata (non-visible)
        self._add_security_metadata(doc)

        # Embed payloads using stealth techniques
        techniques_used = []
        embedded_payloads = []

        if payload_embeddings:
            for embedding_config in payload_embeddings:
                try:
                    technique = embedding_config.get("technique", "font_color_match")
                    payload = embedding_config.get("payload", "")
                    position = embedding_config.get("position", "end")
                    anchor_text = embedding_config.get("anchor_text")
                    additional_config = embedding_config.get("additional_config", {})

                    self._embed_payload(
                        doc,
                        payload,
                        technique,
                        position,
                        anchor_text,
                        additional_config,
                    )

                    techniques_used.append(technique)
                    embedded_payloads.append(payload[:100] + "..." if len(payload) > 100 else payload)

                except Exception as e:
                    logger.warning(f"Failed to embed payload: {e}")
                    # Continue with other payloads

        # Save document to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        file_content = doc_bytes.read()
        file_size = len(file_content)

        # Generate filename if not provided
        if not filename:
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}"

        # Prepare response
        result = {
            "filename": f"{filename}.docx",
            "file_size_bytes": file_size,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "template_used": style.value,
            "embedded_payloads_count": len(embedded_payloads),
            "stealth_techniques_used": techniques_used,
            "generated_at": datetime.utcnow().isoformat(),
        }

        if return_base64:
            result["content_base64"] = base64.b64encode(file_content).decode("ascii")

        logger.info(
            f"Generated DOCX document: {filename}.docx "
            f"({file_size} bytes, {len(embedded_payloads)} payloads)"
        )

        return result

    def _apply_template(
        self,
        doc: "Document",
        style: DocumentStyle,
        content: str,
    ) -> None:
        """Apply document template structure."""
        template_info = self.DOCUMENT_TEMPLATES.get(style, self.DOCUMENT_TEMPLATES[DocumentStyle.BUSINESS_MEMO])

        # Add title if applicable
        if template_info.get("title"):
            title = doc.add_heading(template_info["title"], 0)
            title.alignment = 0  # Left alignment

        # Add header fields for structured documents
        if template_info.get("headers"):
            for header in template_info["headers"]:
                para = doc.add_paragraph()
                para.add_run(header).bold = True

            doc.add_paragraph()  # Spacing

        # Add visible content
        content_paragraphs = content.split("\n\n")
        for para_text in content_paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

    def _apply_metadata(
        self,
        doc: "Document",
        metadata: Dict[str, str],
    ) -> None:
        """Apply document metadata."""
        core_props = doc.core_properties

        if "title" in metadata:
            core_props.title = metadata["title"]
        if "author" in metadata:
            core_props.author = metadata["author"]
        if "subject" in metadata:
            core_props.subject = metadata["subject"]
        if "comments" in metadata:
            core_props.comments = metadata["comments"]
        if "keywords" in metadata:
            core_props.keywords = metadata["keywords"]
        if "category" in metadata:
            core_props.category = metadata["category"]

    def _add_security_metadata(self, doc: "Document") -> None:
        """Add security metadata indicating test nature."""
        core_props = doc.core_properties
        security_note = "SECURITY TEST DOCUMENT - For authorized testing only"

        if not core_props.comments:
            core_props.comments = security_note
        else:
            core_props.comments += f" | {security_note}"

        core_props.category = "Security Testing"

    def _embed_payload(
        self,
        doc: "Document",
        payload: str,
        technique: str,
        position: str,
        anchor_text: Optional[str],
        additional_config: Dict[str, Any],
    ) -> None:
        """Embed a payload using the specified stealth technique."""
        from docx.shared import RGBColor, Pt
        from docx.oxml import OxmlElement

        # Find insertion point
        paragraph = self._find_insertion_paragraph(doc, position, anchor_text)

        if technique == "font_color_match":
            self._embed_font_color_match(paragraph, payload, additional_config)

        elif technique == "white_on_white":
            self._embed_white_on_white(paragraph, payload, additional_config)

        elif technique == "tiny_font":
            self._embed_tiny_font(paragraph, payload, additional_config)

        elif technique == "zero_width_unicode":
            self._embed_zero_width_unicode(paragraph, payload, additional_config)

        elif technique == "hidden_text_property":
            self._embed_hidden_property(paragraph, payload, additional_config)

        elif technique == "metadata_injection":
            self._embed_in_metadata(doc, payload, additional_config)

        elif technique == "footnote_embedding":
            self._embed_in_footnote(doc, paragraph, payload, additional_config)

        elif technique == "comment_injection":
            self._embed_in_comment(doc, paragraph, payload, additional_config)

        elif technique == "header_footer_hidden":
            self._embed_in_header_footer(doc, payload, additional_config)

        elif technique == "invisible_table":
            self._embed_in_invisible_table(doc, payload, additional_config)

        elif technique == "overlay_text":
            self._embed_overlay_text(paragraph, payload, additional_config)

    def _find_insertion_paragraph(
        self,
        doc: "Document",
        position: str,
        anchor_text: Optional[str],
    ) -> Any:
        """Find the paragraph where payload should be inserted."""
        paragraphs = doc.paragraphs

        if position == "start":
            if paragraphs:
                return paragraphs[0]
            return doc.add_paragraph()

        elif position == "end":
            if paragraphs:
                return paragraphs[-1]
            return doc.add_paragraph()

        elif position in ("after_text", "before_text") and anchor_text:
            for para in paragraphs:
                if anchor_text in para.text:
                    if position == "after_text":
                        return para
                    else:
                        # For before_text, find the previous paragraph
                        idx = paragraphs.index(para)
                        if idx > 0:
                            return paragraphs[idx - 1]
                        return para
            # If anchor not found, append to end
            if paragraphs:
                return paragraphs[-1]

        # Default: last paragraph
        if paragraphs:
            return paragraphs[-1]
        return doc.add_paragraph()

    def _embed_font_color_match(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload with font color matching background (white on white)."""
        from docx.shared import RGBColor

        # Get custom colors from config or default to white
        font_color_rgb = config.get("font_color", [255, 255, 255])
        if isinstance(font_color_rgb, str):
            # Parse hex color
            font_color_rgb = [
                int(font_color_rgb[i:i+2], 16) for i in (0, 2, 4)
            ]

        run = paragraph.add_run(payload)
        run.font.color.rgb = RGBColor(*font_color_rgb)
        run.font.size = None  # Keep normal size

    def _embed_white_on_white(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload as white text on white background."""
        from docx.shared import RGBColor

        run = paragraph.add_run(payload)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = None

        # Add white highlight if supported
        try:
            from docx.oxml import OxmlElement
            rPr = run._element.get_or_add_rPr()
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'white')
            rPr.append(highlight)
        except Exception:
            pass  # Highlight not critical

    def _embed_tiny_font(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload with tiny font size (1pt)."""
        from docx.shared import Pt

        font_size = config.get("font_size", 1)
        run = paragraph.add_run(payload)
        run.font.size = Pt(font_size)

        # Optionally make it white too
        if config.get("also_hide_color", True):
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(255, 255, 255)

    def _embed_zero_width_unicode(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload using zero-width Unicode characters.

        This encodes the payload as zero-width characters mixed with
        visible content, making it completely invisible to humans
        but readable to machines.
        """
        # Use binary encoding with zero-width characters
        zero_chars = config.get(
            "zero_width_chars",
            ["\u200B", "\u200C"]  # zero-width space, zero-width non-joiner
        )

        encoded = self._encode_as_zero_width(payload, zero_chars)

        # Insert at the beginning of the paragraph
        if paragraph.runs:
            first_run = paragraph.runs[0]
            original_text = first_run.text
            first_run.text = encoded + original_text
        else:
            paragraph.add_run(encoded)

    def _encode_as_zero_width(self, text: str, zero_chars: List[str]) -> str:
        """Encode text as zero-width characters.

        Uses binary encoding: 0 = first zero-width char, 1 = second.
        """
        if len(zero_chars) < 2:
            zero_chars = ["\u200B", "\u200C"]

        zero_0 = zero_chars[0]
        zero_1 = zero_chars[1]

        # Convert text to binary, then to zero-width characters
        binary_str = "".join(format(ord(c), '016b') for c in text)
        encoded = ""

        for bit in binary_str:
            if bit == '0':
                encoded += zero_0
            else:
                encoded += zero_1

        return encoded

    def _embed_hidden_property(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload using the hidden text property (vanish)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run(payload)

        # Set vanish property (hidden text)
        rPr = run._element.get_or_add_rPr()
        vanish = OxmlElement('w:vanish')
        rPr.append(vanish)

    def _embed_in_metadata(
        self,
        doc: "Document",
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload in document metadata."""
        core_props = doc.core_properties
        field = config.get("metadata_field", "comments")

        existing = getattr(core_props, field, "") or ""
        if existing:
            payload = existing + " | " + payload

        setattr(core_props, field, payload)

    def _embed_in_footnote(
        self,
        doc: "Document",
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload as a hidden footnote."""
        # Add a run with a marker
        run = paragraph.add_run("†")
        run.font.superscript = True
        run.font.color.rgb = None  # Visible marker

        # Add the footnote content at the end (hidden)
        # Since python-docx has limited footnote support,
        # we add it as a regular paragraph with tiny white text
        from docx.shared import Pt, RGBColor

        footnote_para = doc.add_paragraph()
        footnote_run = footnote_para.add_run(f"† {payload}")
        footnote_run.font.size = Pt(1)
        footnote_run.font.color.rgb = RGBColor(255, 255, 255)

    def _embed_in_comment(
        self,
        doc: "Document",
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload as a document comment.

        Note: Full comment support requires OXML manipulation.
        This is a simplified implementation.
        """
        # Add a comment marker
        run = paragraph.add_run("[1]")
        run.font.color.rgb = None
        run.font.superscript = True

        # Embed the comment content as hidden text
        from docx.shared import Pt, RGBColor

        comment_para = doc.add_paragraph()
        comment_run = comment_para.add_run(f"[1] {payload}")
        comment_run.font.size = Pt(1)
        comment_run.font.color.rgb = RGBColor(255, 255, 255)

    def _embed_in_header_footer(
        self,
        doc: "Document",
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload in document header or footer (hidden)."""
        from docx.shared import Pt, RGBColor

        section = doc.sections[0]

        location = config.get("location", "footer")

        if location == "header":
            container = section.header
        else:
            container = section.footer

        # Clear existing content
        for para in container.paragraphs:
            para.clear()

        # Add hidden payload
        para = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        run = para.add_run(payload)
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(255, 255, 255)

    def _embed_in_invisible_table(
        self,
        doc: "Document",
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload in an invisible table."""
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Add table
        table = doc.add_table(rows=1, cols=1)
        table.alignment = 0  # Left

        cell = table.cell(0, 0)
        para = cell.paragraphs[0]
        run = para.add_run(payload)
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(255, 255, 255)

        # Remove table borders
        tbl = table._tbl
        tblPr = tbl.tblPr

        # Get or create tblBorders
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)

    def _embed_overlay_text(
        self,
        paragraph: Any,
        payload: str,
        config: Dict[str, Any],
    ) -> None:
        """Embed payload as overlay text (same text, hidden)."""
        from docx.shared import Pt, RGBColor

        # Get the existing text in the paragraph
        if paragraph.runs:
            existing_text = paragraph.text

            # Add the same text as invisible overlay
            run = paragraph.add_run(existing_text)
            run.font.size = Pt(1)
            run.font.color.rgb = RGBColor(255, 255, 255)

            # Add the payload as well
            payload_run = paragraph.add_run(payload)
            payload_run.font.size = Pt(1)
            payload_run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            # No existing text, just add hidden payload
            run = paragraph.add_run(payload)
            run.font.size = Pt(1)
            run.font.color.rgb = RGBColor(255, 255, 255)


# Update ToolFactory to include DocumentGeneratorTool
def register_document_generator_tool() -> None:
    """Register the DocumentGeneratorTool with the ToolFactory."""
    from . import ToolFactory
    if "docx_generator" not in ToolFactory._instances:
        ToolFactory._instances["docx_generator"] = DocumentGeneratorTool()


# Auto-register on import
try:
    register_document_generator_tool()
except Exception:
    pass  # May fail if imported during initial module load
