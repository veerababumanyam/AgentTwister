"""
Shared Tools Library

Declarative tools that can be bound to agents for common operations:
- Database readers/writers (SQLite)
- File parsers (PDF, DOCX, TXT, JSON)
- HTTP callers with retry logic
"""

import asyncio
import json
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from enum import Enum
from io import BytesIO
from typing import Any, AsyncIterator, Dict, List, Optional, Union

import httpx

from ..base_agent import ToolDefinition
from ..resilience import ExponentialBackoff, CircuitBreaker, CircuitBreakerConfig

logger = logging.getLogger(__name__)


# ============================================================
# BASE TOOL CLASS
# ============================================================

class BaseTool(ABC):
    """
    Base class for all tools.

    Tools encapsulate specific capabilities (file parsing, HTTP calls,
    database operations) that can be bound to agents.
    """

    def __init__(self, name: str, description: str, is_dangerous: bool = False):
        self.name = name
        self.description = description
        self.is_dangerous = is_dangerous

    @abstractmethod
    async def call(self, **kwargs) -> Any:
        """Execute the tool."""
        pass

    def to_definition(self) -> ToolDefinition:
        """Convert to ToolDefinition for agent registration."""
        return ToolDefinition(
            name=self.name,
            description=self.description,
            parameters=self.get_parameters_schema(),
            handler=self.call,
            is_dangerous=self.is_dangerous,
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        """Get JSON schema for tool parameters."""
        return {"type": "object"}


# ============================================================
# DATABASE TOOLS (SQLite)
# ============================================================

class DatabaseReaderTool(BaseTool):
    """
    Tool for reading from SQLite database.

    Supports:
    - Read records by ID
    - Query with filters
    - List records from table
    """

    def __init__(self):
        super().__init__(
            name="database_read",
            description="Read records from SQLite database. Supports queries and filtering.",
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "table": {
                    "type": "string",
                    "description": "Table name to read from",
                },
                "record_id": {
                    "type": "string",
                    "description": "Record ID (optional, if not provided performs query/list)",
                },
                "filters": {
                    "type": "object",
                    "description": "Filters for query (field: value pairs)",
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum number of results (default: 100)",
                },
            },
            "required": ["table"],
        }

    async def call(
        self,
        table: str,
        record_id: Optional[str] = None,
        filters: Optional[Dict[str, Any]] = None,
        limit: int = 100,
    ) -> Union[Dict[str, Any], List[Dict[str, Any]]]:
        """
        Read from database.

        Args:
            table: Table name
            record_id: Specific record ID to read
            filters: Query filters
            limit: Maximum results

        Returns:
            Record dict or list of records
        """
        from ..database import async_session_maker
        from sqlalchemy import text

        async with async_session_maker() as session:
            if record_id:
                # Read single record
                query = text(f"SELECT * FROM {table} WHERE id = :id")
                result = await session.execute(query, {"id": record_id})
                row = result.fetchone()
                if row:
                    result_dict = dict(row._mapping)
                    logger.debug(f"Database read: {table}/{record_id}")
                    return result_dict
                else:
                    logger.warning(f"Database record not found: {table}/{record_id}")
                    return None

            # Query or list table
            query_str = f"SELECT * FROM {table}"
            params = {}

            # Apply filters
            if filters:
                conditions = []
                for i, (field, value) in enumerate(filters.items()):
                    conditions.append(f"{field} = :filter_{i}")
                    params[f"filter_{i}"] = value
                query_str += " WHERE " + " AND ".join(conditions)

            query_str += f" LIMIT {limit}"
            query = text(query_str)

            result = await session.execute(query, params)
            rows = result.fetchall()
            results = [dict(row._mapping) for row in rows]

            logger.debug(f"Database query: {table} returned {len(results)} results")
            return results


class DatabaseWriterTool(BaseTool):
    """
    Tool for writing to SQLite database.

    Supports:
    - Create new record
    - Update existing record
    - Delete record
    """

    def __init__(self):
        super().__init__(
            name="database_write",
            description="Write to SQLite database. Supports create, update, and delete operations.",
            is_dangerous=True,  # Can modify data
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "table": {
                    "type": "string",
                    "description": "Table name to write to",
                },
                "record_id": {
                    "type": "string",
                    "description": "Record ID (for update/delete)",
                },
                "data": {
                    "type": "object",
                    "description": "Data to write (for create/update)",
                },
                "operation": {
                    "type": "string",
                    "enum": ["create", "update", "delete"],
                    "description": "Operation type",
                },
            },
            "required": ["table", "operation"],
        }

    async def call(
        self,
        table: str,
        operation: str,
        record_id: Optional[str] = None,
        data: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Write to database.

        Args:
            table: Table name
            operation: Operation type (create, update, delete)
            record_id: Record ID
            data: Data to write

        Returns:
            Result dict with record ID and operation status
        """
        from ..database import async_session_maker
        from sqlalchemy import text
        import uuid

        async with async_session_maker() as session:
            if operation == "delete":
                if not record_id:
                    raise ValueError("record_id required for delete operation")
                query = text(f"DELETE FROM {table} WHERE id = :id")
                await session.execute(query, {"id": record_id})
                await session.commit()
                logger.info(f"Database delete: {table}/{record_id}")
                return {"id": record_id, "operation": "deleted"}

            # For create/update, we need data
            if data is None:
                raise ValueError("data required for create/update operations")

            if operation == "create":
                record_id = record_id or str(uuid.uuid4())
                data["id"] = record_id
                columns = ", ".join(data.keys())
                placeholders = ", ".join([f":{k}" for k in data.keys()])
                query = text(f"INSERT INTO {table} ({columns}) VALUES ({placeholders})")
                await session.execute(query, data)
                await session.commit()
                logger.info(f"Database create: {table}/{record_id}")
                return {"id": record_id, "operation": "created"}

            elif operation == "update":
                if not record_id:
                    raise ValueError("record_id required for update operation")
                set_clause = ", ".join([f"{k} = :{k}" for k in data.keys()])
                data["id"] = record_id
                query = text(f"UPDATE {table} SET {set_clause} WHERE id = :id")
                await session.execute(query, data)
                await session.commit()
                logger.info(f"Database update: {table}/{record_id}")
                return {"id": record_id, "operation": "updated"}

            else:
                raise ValueError(f"Unknown operation: {operation}")


# ============================================================
# FILE PARSER TOOLS
# ============================================================

class FileFormat(Enum):
    """Supported file formats."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    JSON = "json"
    MD = "md"


class FileParserTool(BaseTool):
    """
    Tool for parsing various file formats.

    Supports:
    - PDF text extraction
    - DOCX text extraction and metadata
    - Plain text files
    - JSON parsing
    - Markdown files
    """

    def __init__(self):
        super().__init__(
            name="file_parser",
            description="Parse and extract content from files (PDF, DOCX, TXT, JSON, MD)",
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to file to parse",
                },
                "file_content": {
                    "type": "string",
                    "description": "Base64-encoded file content (alternative to file_path)",
                },
                "format": {
                    "type": "string",
                    "enum": ["pdf", "docx", "txt", "json", "md", "auto"],
                    "description": "File format (auto-detect if not specified)",
                },
                "extract_metadata": {
                    "type": "boolean",
                    "description": "Extract metadata if available",
                },
            },
        }

    async def call(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[str] = None,
        format: str = "auto",
        extract_metadata: bool = False,
    ) -> Dict[str, Any]:
        """Parse file and extract content."""
        import base64
        import os

        # Determine content source
        if file_path:
            with open(file_path, "rb") as f:
                content_bytes = f.read()
            detected_format = self._detect_format(file_path)
        elif file_content:
            content_bytes = base64.b64decode(file_content)
            detected_format = None
        else:
            raise ValueError("Either file_path or file_content must be provided")

        # Determine format
        if format == "auto":
            format = detected_format or self._detect_format_from_bytes(content_bytes)

        # Parse based on format
        result = {"text": "", "metadata": {}, "format": format}

        if format == "pdf":
            result.update(await self._parse_pdf(content_bytes, extract_metadata))
        elif format == "docx":
            result.update(await self._parse_docx(content_bytes, extract_metadata))
        elif format == "json":
            result.update(await self._parse_json(content_bytes))
        elif format in ("txt", "md"):
            result["text"] = content_bytes.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported format: {format}")

        logger.debug(f"FileParser parsed {format} file: {len(result['text'])} chars")
        return result

    def _detect_format(self, filename: str) -> str:
        """Detect format from filename extension."""
        import os
        _, ext = os.path.splitext(filename.lower())
        return ext.lstrip(".")

    def _detect_format_from_bytes(self, content: bytes) -> str:
        """Detect format from file content bytes."""
        if content.startswith(b"%PDF"):
            return "pdf"
        if content.startswith(b"PK\x03\x04"):
            return "docx"
        stripped = content.strip()
        if stripped.startswith(b"{") or stripped.startswith(b"["):
            return "json"
        return "txt"

    async def _parse_pdf(self, content: bytes, extract_metadata: bool) -> Dict[str, Any]:
        """Parse PDF file."""
        try:
            import PyPDF2
            from io import BytesIO

            pdf_file = BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())

            result = {"text": "\n\n".join(text_parts)}

            if extract_metadata:
                metadata = reader.metadata
                if metadata:
                    result["metadata"] = {
                        "title": metadata.get("/Title", ""),
                        "author": metadata.get("/Author", ""),
                        "page_count": len(reader.pages),
                    }

            return result

        except ImportError:
            logger.error("No PDF library available. Install PyPDF2.")
            raise ImportError("PDF parsing requires PyPDF2")

    async def _parse_docx(self, content: bytes, extract_metadata: bool) -> Dict[str, Any]:
        """Parse DOCX file."""
        try:
            from docx import Document
            from io import BytesIO

            doc_file = BytesIO(content)
            doc = Document(doc_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            result = {"text": "\n\n".join(text_parts)}

            if extract_metadata:
                core_props = doc.core_properties
                result["metadata"] = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "paragraph_count": len(doc.paragraphs),
                }

            return result

        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("DOCX parsing requires python-docx")

    async def _parse_json(self, content: bytes) -> Dict[str, Any]:
        """Parse JSON file."""
        text = content.decode("utf-8")
        data = json.loads(text)
        return {
            "text": json.dumps(data, indent=2),
            "metadata": {"parsed": True},
            "data": data,
        }


# ============================================================
# HTTP CALLER TOOL
# ============================================================

class HTTPCallerTool(BaseTool):
    """
    Tool for making HTTP requests with retry and circuit breaker.

    Features:
    - Automatic retry with exponential backoff
    - Circuit breaker for failing endpoints
    - Streaming support
    - Request/response logging
    """

    def __init__(
        self,
        max_retries: int = 3,
        enable_circuit_breaker: bool = True,
    ):
        super().__init__(
            name="http_caller",
            description="Make HTTP requests with retry logic and circuit breaker protection.",
            is_dangerous=True,  # Can make external requests
        )
        self._max_retries = max_retries
        self._enable_circuit_breaker = enable_circuit_breaker
        self._circuit_breakers: Dict[str, CircuitBreaker] = {}
        self._client = httpx.AsyncClient(
            timeout=httpx.Timeout(120.0),
            limits=httpx.Limits(max_keepalive_connections=20, max_connections=100),
        )

    def get_parameters_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "Request URL"},
                "method": {
                    "type": "string",
                    "enum": ["GET", "POST", "PUT", "PATCH", "DELETE"],
                    "description": "HTTP method",
                },
                "headers": {"type": "object", "description": "Request headers"},
                "json_body": {"type": "object", "description": "JSON request body"},
                "params": {"type": "object", "description": "Query parameters"},
                "timeout": {"type": "number", "description": "Request timeout in seconds"},
            },
            "required": ["url", "method"],
        }

    async def call(
        self,
        url: str,
        method: str = "GET",
        headers: Optional[Dict[str, str]] = None,
        json_body: Optional[Dict[str, Any]] = None,
        params: Optional[Dict[str, str]] = None,
        timeout: Optional[float] = None,
    ) -> Dict[str, Any]:
        """Make HTTP request."""
        method = method.upper()
        base_url = self._get_base_url(url)
        breaker = self._get_circuit_breaker(base_url)

        async def _make_request() -> Dict[str, Any]:
            request_headers = {
                "User-Agent": "AgentTwister/1.0",
                **(headers or {}),
            }

            request_kwargs = {
                "method": method,
                "url": url,
                "headers": request_headers,
                "params": params,
                "timeout": timeout or 120.0,
            }

            if json_body and method in ("POST", "PUT", "PATCH"):
                request_kwargs["json"] = json_body

            logger.debug(f"HTTP {method} {url}")
            response = await self._client.request(**request_kwargs)

            try:
                body = response.json()
            except ValueError:
                body = response.text

            logger.debug(f"HTTP {method} {url} -> {response.status_code}")

            return {
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "body": body,
                "text": response.text,
                "success": 200 <= response.status_code < 300,
            }

        # Use circuit breaker if enabled
        if self._enable_circuit_breaker:
            try:
                return await breaker.call(_make_request)
            except Exception as e:
                if "CircuitBreakerOpenError" in str(type(e)):
                    raise

        # Retry with backoff
        backoff = ExponentialBackoff(max_attempts=self._max_retries)
        last_error = None

        async for attempt in backoff:
            try:
                return await _make_request()
            except Exception as e:
                last_error = e
                logger.warning(f"HTTP request attempt {attempt + 1} failed: {e}")

        raise last_error or Exception("HTTP request failed")

    async def call_stream(
        self,
        url: str,
        json_body: Optional[Dict[str, Any]] = None,
        method: str = "POST",
        headers: Optional[Dict[str, str]] = None,
    ) -> AsyncIterator[str]:
        """Make streaming HTTP request."""
        request_headers = {
            "User-Agent": "AgentTwister/1.0",
            **(headers or {}),
        }

        async with self._client.stream(
            method=method,
            url=url,
            json=json_body,
            headers=request_headers,
            timeout=120.0,
        ) as response:
            response.raise_for_status()

            async for chunk in response.aiter_bytes():
                if chunk:
                    chunk_text = chunk.decode("utf-8", errors="replace")
                    for line in chunk_text.split("\n"):
                        if line.startswith("data: "):
                            data = line[6:]
                            if data.strip() == "[DONE]":
                                return
                            try:
                                parsed = json.loads(data)
                                if "choices" in parsed:
                                    delta = parsed["choices"][0].get("delta", {})
                                    if "content" in delta:
                                        yield delta["content"]
                            except json.JSONDecodeError:
                                if data.strip():
                                    yield data
                        elif line.strip() and not line.startswith(":"):
                            yield line

    def _get_base_url(self, url: str) -> str:
        from urllib.parse import urlparse
        parsed = urlparse(url)
        return f"{parsed.scheme}://{parsed.netloc}"

    def _get_circuit_breaker(self, base_url: str) -> CircuitBreaker:
        if base_url not in self._circuit_breakers:
            self._circuit_breakers[base_url] = CircuitBreaker(
                name=f"http_{base_url}",
                config=CircuitBreakerConfig(failure_threshold=5),
            )
        return self._circuit_breakers[base_url]

    async def close(self) -> None:
        await self._client.aclose()


# ============================================================
# TOOL FACTORY
# ============================================================

class ToolFactory:
    """Factory for creating tool instances."""

    _instances: Dict[str, BaseTool] = {}

    @classmethod
    def get_database_reader(cls) -> DatabaseReaderTool:
        if "database_reader" not in cls._instances:
            cls._instances["database_reader"] = DatabaseReaderTool()
        return cls._instances["database_reader"]

    @classmethod
    def get_database_writer(cls) -> DatabaseWriterTool:
        if "database_writer" not in cls._instances:
            cls._instances["database_writer"] = DatabaseWriterTool()
        return cls._instances["database_writer"]

    @classmethod
    def get_file_parser(cls) -> FileParserTool:
        if "file_parser" not in cls._instances:
            cls._instances["file_parser"] = FileParserTool()
        return cls._instances["file_parser"]

    @classmethod
    def get_http_caller(cls) -> HTTPCallerTool:
        if "http_caller" not in cls._instances:
            cls._instances["http_caller"] = HTTPCallerTool()
        return cls._instances["http_caller"]

    @classmethod
    def get_docx_generator(cls) -> "DocumentGeneratorTool":
        from .docx_generator import DocumentGeneratorTool
        if "docx_generator" not in cls._instances:
            cls._instances["docx_generator"] = DocumentGeneratorTool()
        return cls._instances["docx_generator"]

    @classmethod
    def get_pdf_generator(cls) -> "PDFGeneratorTool":
        from .pdf_generator import PDFGeneratorTool
        if "pdf_generator" not in cls._instances:
            cls._instances["pdf_generator"] = PDFGeneratorTool()
        return cls._instances["pdf_generator"]

    # Backward compatibility aliases
    get_firestore_reader = get_database_reader
    get_firestore_writer = get_database_writer

    @classmethod
    async def close_all(cls) -> None:
        if "http_caller" in cls._instances:
            await cls._instances["http_caller"].close()
        cls._instances.clear()


# Aliases for backward compatibility
FirestoreReaderTool = DatabaseReaderTool
FirestoreWriterTool = DatabaseWriterTool
